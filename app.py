import streamlit as st
import cv2
import mediapipe as mp
import numpy as np
import time
import queue
import av
import pyttsx3
import threading
from io import BytesIO
from docx import Document
from datetime import datetime
from tensorflow.keras.models import load_model
from streamlit_webrtc import webrtc_streamer, WebRtcMode

from app import result_queue

# Styling
style = """
<style>
    /* Adjust block container padding */
    .reportview-container .main .block-container {
        padding-top: 1rem;
        padding-bottom: 1rem;
        padding-left: 1rem;
        padding-right: 1rem;
    }
</style>
<style>
    #MainMenu{visibility: hidden;}
</style>
"""
st.set_page_config(page_title='Hand Gesture Recognition')
st.markdown(style, unsafe_allow_html=True)
st.title("Hand Gesture Recognition")

# Load model
@st.cache(allow_output_mutation=True)
def load_recognition_model():
    return load_model("./model/model_fold_1.keras")
model = load_recognition_model()

#mediapipe & prediction setup
mp_hands = mp.solutions.hands
mp_drawing = mp.solutions.drawing_utils
mp_drawing_styles = mp.solutions.drawing_styles

D_FEATURES = 20
N_FEATURES = 1
CONFIDENCE_THRESHOLD = 0.7

labels_dict = {
    0: 'A', 1: 'B', 2: 'C', 3: 'D', 4: 'E', 5: 'F',
    6: 'G', 7: 'H', 8: 'I', 9: 'J', 10: 'K', 11: 'L',
    12: 'M', 13: 'N', 14: 'O', 15: 'P', 16: 'Q', 17: 'R',
    18: 'S', 19: 'T', 20: 'U', 21: 'V', 22: 'W', 23: 'X',
    24: 'Y', 25: 'Z', 26: 'del', 27: 'space'
}
result_queue = queue.Queue()
# Session-state defaults
# if "tts_engine" not in st.session_state:
    # st.session_state.tts_engine = pyttsx3.init()
defaults = {
    "detected_word": "",
    "last_char":     None,
    "hold_start":    0.0,
    "last_print":    0.0,
    "paused":        False,
    "spoken_word":   None,
    "tts_thread":    None,
    "tts_engine":   None,
    "cooldown":      1.0,   # user‐adjustable later
}
for k,v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v

def speak_text_async(text):
    engine = st.session_state.tts_engine
    prev = st.session_state.tts_thread
    # """Non‐blocking TTS; only speaks when invoked by button."""
    if prev and prev.is_alive():
        engine.stop()
        prev.join()

    def _run():
        # create a fresh engine each time
        engine = pyttsx3.init()
        engine.say(text)
        engine.runAndWait()

    t = threading.Thread(target=_run, daemon=True)
    t.start()
    st.session_state.tts_thread   = t
    st.session_state.spoken_word  = text

def emit_char(ch):
    if ch=='space': st.session_state.detected_word += ' '
    else: st.session_state.detected_word += ch

def delete_one_char():
    """Safely remove exactly one character, if any."""
    w = st.session_state.detected_word or " "
    st.session_state.detected_word = w[:-1]

#Function FRAME PROCESS
def process_frame(frame):
    image = frame.to_ndarray(format="bgr24")
    image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
    predicted_character = None
    data_aux = []

    with mp_hands.Hands(
        static_image_mode=False,
        max_num_hands=1,
        min_detection_confidence=0.5,
        min_tracking_confidence=0.5
    ) as hands:
        results = hands.process(image_rgb)

        list_knuckle = [
            mp_hands.HandLandmark.WRIST,
            mp_hands.HandLandmark.THUMB_TIP,
            mp_hands.HandLandmark.INDEX_FINGER_TIP,
            mp_hands.HandLandmark.MIDDLE_FINGER_TIP,
            mp_hands.HandLandmark.RING_FINGER_TIP,
            mp_hands.HandLandmark.PINKY_TIP,
            mp_hands.HandLandmark.MIDDLE_FINGER_MCP,
            mp_hands.HandLandmark.RING_FINGER_MCP,
            mp_hands.HandLandmark.PINKY_MCP,
            mp_hands.HandLandmark.INDEX_FINGER_MCP
        ]

        if results.multi_hand_landmarks:
            for hand_landmarks in results.multi_hand_landmarks:
                mp_drawing.draw_landmarks(
                    image, hand_landmarks, mp_hands.HAND_CONNECTIONS,
                    mp_drawing_styles.get_default_hand_landmarks_style(),
                    mp_drawing_styles.get_default_hand_connections_style()
                )
                for i in list_knuckle:
                    x = hand_landmarks.landmark[i].x
                    y = hand_landmarks.landmark[i].y
                    data_aux.append(x)
                    data_aux.append(y)

                if len(data_aux) == 20:
                    X_test = np.asarray([np.asarray(data_aux)], dtype=np.float32).reshape(-1, D_FEATURES, N_FEATURES)
                    prediction = model.predict(X_test, verbose=0)[0]
                    confidence = np.max(prediction)
                    if confidence >= CONFIDENCE_THRESHOLD:
                        predicted_character = labels_dict.get(np.argmax(prediction))
                        cv2.putText(image, f"Predicted: {predicted_character}", (10, 30),
                                    cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2, cv2.LINE_AA)

    result_queue.put(predicted_character)
    return av.VideoFrame.from_ndarray(image, format="bgr24")


f_col1, f_col2 = st.columns([1,1], gap="large")
with f_col1:
    #Using device's camera
    st.write("Show a hand gesture to the camera to get real-time predictions")
    # result_queue = queue.Queue()
    webrtc_ctx = webrtc_streamer(
        key = "hand-gesture",
        mode = WebRtcMode.SENDRECV,
        rtc_configuration = {"iceServers": [
        {"urls": "stun:stun.l.google.com:19302"},
        # {"urls": "turn:global.relay.metered.ca:80", "username": st.secrets['metered_username'], "credential": st.secrets['metered_credentials']},
        # {"urls": "turn:global.relay.metered.ca:80?transport=tcp", "username": st.secrets['metered_username'], "credential": st.secrets['metered_credentials']},
        # {"urls": "turn:global.relay.metered.ca:443", "username": st.secrets['metered_username'], "credential": st.secrets['metered_credentials']},
        # {"urls": "turns:global.relay.metered.ca:443?transport=tcp", "username": st.secrets['metered_username'], "credential": st.secrets['metered_credentials']},
    ]},
        video_frame_callback = process_frame,
        media_stream_constraints = {"video": True, "audio": False},
        async_processing = True,
        # reuse_video_stream=True,
    )
with f_col2:
    st.image("signlanguage.jpg", caption="The Alphabet of Sign Language", use_column_width=True)

d_col1, d_col2 = st.columns([2,1],gap = 'small')
with d_col1:
    if st.button('⏸️ Pause',key = 'pause_btn'):
        st.session_state.paused = True
    # else: st.session_state.paused=False
with d_col2:
    if st.button('▶️ Resume',key = 'resume_btn'):
        st.session_state.paused = False
st.session_state.cooldown = st.slider('Cooldown for repeat (s)',0.5,3.0,st.session_state.cooldown,0.05)

# Create placeholders
pred_ph = st.empty()
word_ph = st.empty()

c1, c2, c3, c4 = st.columns(4, gap = 'small')
with c1:
    if st.button('Clear', key = 'btn_clear'):
        delete_one_char()
with c2:
    if st.button('🗑️ Reset', key = 'btn_remove'):
        st.session_state.detected_word = ' '
with c3:
    if st.button('🔊 Speak', key = 'btn_speak'):
        speak_text_async(st.session_state.detected_word)
with c4:
    buf = BytesIO()
    doc = Document()
    ts = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    doc.add_heading(f'Recognized Word – {ts}', level=1)
    doc.add_paragraph(st.session_state.detected_word)
    doc.save(buf);
    buf.seek(0)
    download_clicked = c4.download_button(
        '⬇️ Download .docx', data=buf,
        file_name='recognized_word.docx',
        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        key='btn_download'
    )
# speak_btn = st.button('🔊 Speak', key = 'btn_speak')
# Prediction Loop
if webrtc_ctx.state.playing:
    while webrtc_ctx.state.playing:
        try:
            ch = result_queue.get(timeout=1)
        except queue.Empty:
            pred_ph.info('⏳ Waiting for prediction…')
            continue

        now = time.time()

        #Outer pause guard
        if ch and not st.session_state.paused:
            #Del branch
            if ch == 'del':
                held = now - st.session_state.hold_start
                if ch != st.session_state.last_char:
                    st.session_state.last_char = ch
                    st.session_state.hold_start = now
                    st.session_state.last_print = 0.0
                elif held >= 1.0 and (now - st.session_state.last_print) >= st.session_state.cooldown:
                    delete_one_char()
                    st.session_state.last_print = now
                    st.session_state.hold_start = now
            #Space branch
            elif ch == 'space':
                if ch != st.session_state.last_char:
                    emit_char(' ')
                    st.session_state.last_char = ch
                    st.session_state.hold_start = now
                    st.session_state.last_print = now
            #Letter branch
            else:
                if ch != st.session_state.last_char:
                    emit_char(ch)
                    st.session_state.last_print = now
                    st.session_state.hold_start = now
                elif now - st.session_state.last_print >= 2.0:
                    emit_char(ch)
                    st.session_state.last_print = now

            st.session_state.last_char = ch
        # if speak_btn:
            # speak_text_async(st.session_state.detected_word)

            # Update UI
        pred_ph.success(f"✅ Current: {ch}")
        word_ph.markdown(f"## ✍️ {st.session_state.detected_word}")
            # Short sleep to allow UI rendering
        time.sleep(0.05)



